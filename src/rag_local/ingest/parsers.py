"""문서 파서: PDF, TXT, DOCX 리더.

각 포맷별 파서가 List[Document]를 반환한다.
PDF는 텍스트 추출 우선, 스캔 PDF는 EasyOCR 폴백.
Phase 2에서 HWP/HWPX 파서를 추가할 예정.
"""

from __future__ import annotations

import logging
from pathlib import Path

from llama_index.core.schema import Document

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}

# 페이지당 텍스트가 이 글자수 미만이면 스캔 PDF로 판단
_OCR_THRESHOLD = 50

# EasyOCR Reader 싱글톤
_ocr_reader = None


def _get_ocr_reader():
    """EasyOCR Reader를 싱글톤으로 반환한다."""
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr

        logger.info("EasyOCR Reader 초기화 (ko, en)...")
        _ocr_reader = easyocr.Reader(
            ["ko", "en"],
            gpu=False,  # GPU 없어도 동작
        )
        logger.info("EasyOCR Reader 초기화 완료")
    return _ocr_reader


def _ocr_page(page) -> str:
    """PyMuPDF 페이지를 이미지로 렌더링 후 EasyOCR로 텍스트 추출."""
    import numpy as np

    # 300 DPI로 렌더링
    pix = page.get_pixmap(dpi=300)
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
        pix.height, pix.width, pix.n
    )

    reader = _get_ocr_reader()
    results = reader.readtext(img, detail=0, paragraph=True)
    return "\n".join(results)


def parse_pdf(file_path: Path) -> list[Document]:
    """PDF 파일을 파싱한다.

    텍스트 PDF는 PyMuPDF로 직접 추출.
    스캔 PDF(페이지당 텍스트 50자 미만)는 EasyOCR로 자동 전환.
    """
    import pymupdf

    docs = []
    pdf = pymupdf.open(str(file_path))
    total_pages = len(pdf)
    ocr_pages = 0

    for page_num, page in enumerate(pdf, start=1):
        text = page.get_text()

        # 텍스트가 충분하지 않으면 스캔 페이지로 판단 → OCR
        if len(text.strip()) < _OCR_THRESHOLD:
            logger.info(
                "스캔 페이지 감지: %s (p.%d) — OCR 수행 중...",
                file_path.name, page_num,
            )
            text = _ocr_page(page)
            ocr_pages += 1

        is_ocr = len(page.get_text().strip()) < _OCR_THRESHOLD
        if text.strip():
            docs.append(Document(
                text=text,
                metadata={
                    "file_name": file_path.name,
                    "file_path": str(file_path),
                    "file_type": "pdf",
                    "page_number": page_num,
                    "total_pages": total_pages,
                    "ocr": is_ocr,
                },
            ))

    pdf.close()

    if ocr_pages > 0:
        logger.info(
            "PDF 파싱 완료: %s (%d페이지, OCR: %d페이지)",
            file_path.name, len(docs), ocr_pages,
        )
    else:
        logger.info("PDF 파싱 완료: %s (%d페이지)", file_path.name, len(docs))

    return docs


def parse_txt(file_path: Path) -> list[Document]:
    """TXT 파일을 파싱한다."""
    text = file_path.read_text(encoding="utf-8")
    if not text.strip():
        logger.warning("빈 텍스트 파일: %s", file_path.name)
        return []
    logger.info("TXT 파싱 완료: %s", file_path.name)
    return [Document(
        text=text,
        metadata={
            "file_name": file_path.name,
            "file_path": str(file_path),
            "file_type": "txt",
        },
    )]


def parse_docx(file_path: Path) -> list[Document]:
    """DOCX 파일을 파싱한다 (python-docx)."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    if not text.strip():
        logger.warning("빈 DOCX 파일: %s", file_path.name)
        return []
    logger.info("DOCX 파싱 완료: %s (%d 단락)", file_path.name, len(paragraphs))
    return [Document(
        text=text,
        metadata={
            "file_name": file_path.name,
            "file_path": str(file_path),
            "file_type": "docx",
        },
    )]


def parse_file(file_path: Path) -> list[Document]:
    """파일 확장자에 따라 적절한 파서를 선택하여 파싱한다."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    else:
        logger.warning("지원하지 않는 파일 형식: %s", ext)
        return []
